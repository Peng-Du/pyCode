import os
import docx
import azure.cognitiveservices.speech as speechsdk

# 打开.docx文件
doc = docx.Document("test.docx")

# Set up the Azure Speech API configuration
speech_key = os.environ.get('AZURE_API_KEY')
service_region = 'eastus'

speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)

# 遍历所有段落
i = 0
while i < len(doc.paragraphs):
    # 获取当前段落
    para = doc.paragraphs[i]
    # 如果段落以“P+数字”格式命名
    if para.text.startswith("P") and para.text[1:].isdigit():
        # 输出段落名
        # print(f"段落名：{para.text}")
        chaptername = para.text
        # 输出段落内容
        i += 1
        content = ""
        while i < len(doc.paragraphs):
            next_para = doc.paragraphs[i]
            # 如果下一个段落也以“P+数字”格式命名，说明当前段落内容已经结束
            if next_para.text.startswith("P") and next_para.text[1:].isdigit():
                break
            # 否则将当前段落内容加入到content中
            content += next_para.text
            content += "\n"  # 加入回车
            i += 1
        # print(f"内容：{content}")

        #设置输出文件
        audio_config = speechsdk.audio.AudioOutputConfig(filename=chaptername+".wav")

        # Create a synthesizer using the configured settings
        speech_synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config, audio_config=audio_config)

        # Define the ssml to be synthesized
        ssml = """
        <speak version="1.0" xmlns="http://www.w3.org/2001/10/synthesis" xml:lang="en-US">
            <voice name="es-ES-ElviraNeural">
                <prosody rate="0.9">""" + content +"""
                </prosody>
            </voice>
        </speak>
        """
        # Synthesize the ssml to an MP3 file
        result = speech_synthesizer.speak_ssml_async(ssml).get()


    else:
        i += 1
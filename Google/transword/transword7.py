from google.cloud import translate_v3 as translate
import os
import docx

def split_list_into_subsets(contents, subset_size):
    for i in range(0, len(contents), subset_size):
        yield contents[i:i + subset_size]

def translate_word_document(input_file, output_file):
    # 设置Google Cloud API密钥路径
    # os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = ''
    # set GOOGLE_APPLICATION_CREDENTIALS=
    # echo %GOOGLE_APPLICATION_CREDENTIALS%

    # 创建Google Translate客户端
    client = translate.TranslationServiceClient()

    # 加载docx文件
    doc = docx.Document(input_file)

    # 创建翻译请求
    parent = "projects/transppt-123456/locations/global"
    target_language = "es"
    model = "projects/transppt-123456/locations/global/models/general/nmt"
    maxtext = 200
    looptime = 0
    
    while looptime*maxtext < len(doc.paragraphs):
        contents = []
        i = looptime*maxtext
        while i < len(doc.paragraphs) and i < (looptime+1)*maxtext:
            # 获取当前段落
            print('-----------',i,'-----------',doc.paragraphs[i].text)
            text = doc.paragraphs[i].text
            if text:
                contents.append(doc.paragraphs[i].text)
            i = i + 1

        if not contents:
            print("No text content found in the DOCX file. Exiting translation.")
            return

        request = translate.TranslateTextRequest(
            parent=parent,
            target_language_code=target_language,
            contents=contents,
            model=model
        )

        response = client.translate_text(request=request)
        

        # 更新DOCX中的文本
        i = looptime*maxtext
        Index = 0
        while i < len(doc.paragraphs) and i < (looptime+1)*maxtext:
            # 获取当前段落
            text = doc.paragraphs[i].text
            if text:
                doc.paragraphs[i].text = response.translations[Index].translated_text
                Index = Index + 1
            i = i + 1

        looptime = looptime + 1

    # 翻译table中的文本
    totaltable = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    # print('Paragraph Text: cell--',i,'---',paragraph.text)
                    text = paragraph.text
                    if text:
                        totaltable = totaltable + 1

    print('totaltable:',totaltable)

    contents = []
    transcontents = []
    i = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    # print('Paragraph Text: cell--',i,'---',paragraph.text)
                    text = paragraph.text
                    if text:
                        contents.append(text)
                        print('cell--',i,'---',text)
                        i = i + 1

    if not contents:
        print("No text table found in the DOCX file. Exiting translation.")
        # Save the translated document
        doc.save(output_file)
        return

    # print(contents)

    # 将context列表分成多个子集
    subcontents = list(split_list_into_subsets(contents, 1000))

    for i, subset in enumerate(subcontents):
        # print(f"Subset {i + 1} (Length: {len(subset)}):")
        # print(subset)

        request = translate.TranslateTextRequest(
            parent=parent,
            target_language_code=target_language,
            contents=subset,
            model=model
        )

        response = client.translate_text(request=request)

        transcontents.extend(response.translations)

    # print(response)

    i = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    # print('Paragraph Text: cell--',i,'---',paragraph.text)
                    text = paragraph.text
                    if text:
                        paragraph.text = transcontents[i].translated_text
                        print('trans cell--',i,'---',cell.text)
                        i = i + 1

    # Save the translated document
    doc.save(output_file)

# 保存处理后的DOCX文件
input_file = 'test.docx'
output_file = 'test - ES.docx'

# Call the function to translate the Word document
translate_word_document(input_file, output_file)